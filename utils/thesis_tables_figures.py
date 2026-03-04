# fmt: off
"""
Thesis Tables and Figures Generator
====================================
Generates publication-ready tables and figures for the cognitive classification thesis.

This script integrates with the existing analysis pipeline and generates:

TABLES:
- Table 1: Participant characteristics by DEFINITE cognitive status (NCD, MCI, Dementia)
- Table 1b: Participant characteristics by DEFINITE+POSSIBLE cognitive status
- Table 3: Performance metrics for binary classification (NCD vs Dementia, NCD vs MCI)
- Table 4: Performance metrics for three-way classification (multiclass)

FIGURES:
- Figure 1: Forest plot - Balanced accuracy for NCD vs Dementia (includes Jaakkimainen)
- Figure 2: Forest plot - Balanced accuracy for NCD vs MCI
- Figure 3: Forest plot - Balanced accuracy for multiclass classification
- Figure 4: Aggregated feature importance (top 10 features)

Author: Generated for Thesis Analysis
"""

import os
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import seaborn as sns
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

# Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Word document generation disabled.")

# ==============================================================================
# CONFIGURATION
# ==============================================================================

# Paths
DATA_PATH = "MAIN_new.csv"
OUTPUT_BASE_DIR = "cognitive_analysis_output_PARALLEL"
THESIS_OUTPUT_DIR = "thesis_outputs"

# Jaakkimainen algorithm reference values (from literature)
# These are the published performance metrics for the Jaakkimainen dementia case definition
JAAKKIMAINEN_METRICS = {
    'name': 'Jaakkimainen Algorithm',
    'task': '1. Def Normal vs. Def Dementia',
    'balanced_accuracy': 0.74,  # Approximated from Sens/Spec
    'sensitivity': 0.79,
    'specificity': 0.69,
    'ppv': None,  # Not reported
    'npv': None,  # Not reported
    'auc': None,  # Not reported as ROC-AUC
    'reference': 'Jaakkimainen et al., 2016'
}

# Publication style settings
def set_publication_style():
    """Set matplotlib parameters for publication-quality figures (Times New Roman, B&W)."""
    plt.rcParams.update({
        'font.family': 'serif',
        'font.serif': ['Times New Roman', 'Times', 'DejaVu Serif'],
        'font.size': 11,
        'axes.titlesize': 12,
        'axes.labelsize': 11,
        'xtick.labelsize': 10,
        'ytick.labelsize': 10,
        'legend.fontsize': 9,
        'figure.titlesize': 14,
        'axes.linewidth': 1.0,
        'axes.spines.top': False,
        'axes.spines.right': False,
        'axes.edgecolor': 'black',
        'axes.labelcolor': 'black',
        'xtick.color': 'black',
        'ytick.color': 'black',
        'text.color': 'black',
        'figure.dpi': 300,
        'savefig.dpi': 300,
        'savefig.bbox': 'tight',
        'savefig.pad_inches': 0.1,
        'figure.facecolor': 'white',
        'axes.facecolor': 'white',
        'savefig.facecolor': 'white',
    })
    sns.set_style("white")
    sns.set_context("paper", font_scale=1.1)


# ==============================================================================
# TABLE 1: PARTICIPANT CHARACTERISTICS
# ==============================================================================

def calculate_lookback_years(df, visit_date_col='prompt_visitdate'):
    """
    Calculate lookback years available in administrative health data.
    Assumes administrative data starts from approximately 1994 (Alberta health data).
    """
    try:
        df = df.copy()
        df[visit_date_col] = pd.to_datetime(df[visit_date_col], errors='coerce')
        admin_data_start = pd.Timestamp('1994-01-01')
        df['max_lookback_years'] = (df[visit_date_col] - admin_data_start).dt.days / 365.25
        df['max_lookback_years'] = df['max_lookback_years'].clip(lower=0)
        return df['max_lookback_years']
    except Exception as e:
        print(f"Warning: Could not calculate lookback years: {e}")
        return pd.Series([np.nan] * len(df))


def generate_participant_characteristics_table(df, cognitive_status_groups, table_name, output_dir):
    """
    Generate Table 1: Participant characteristics by cognitive status.

    Parameters:
    -----------
    df : DataFrame
        Full dataset
    cognitive_status_groups : dict
        Mapping of group names to cognitive_status values
        e.g., {'NCD': ['definite_normal'], 'MCI': ['definite_mci'], 'Dementia': ['definite_dementia']}
    table_name : str
        Name for the output file
    output_dir : str
        Output directory
    """
    print(f"\n=== Generating {table_name} ===")

    # Prepare data
    df = df.copy()
    df['cognitive_status'] = df['cognitive_status'].fillna('unknown').astype(str)
    df['cognitive_status'] = df['cognitive_status'].str.lower().str.strip().str.replace(' ', '_', regex=False)

    # Calculate lookback years
    df['lookback_years'] = calculate_lookback_years(df)

    # DAD columns (Discharge Abstract Database)
    dad_cols = [c for c in df.columns if c.endswith('_dad')]
    dad_dementia_cols = ['f00_dad', 'f01_dad', 'f02_dad', 'f03_dad', 'g30_dad', 'g310_dad', 'g311_dad']

    # NACRS columns (National Ambulatory Care Reporting System) - treating as part of DAD for this analysis
    nacrs_cols = [c for c in df.columns if c.endswith('_nacrs')]
    nacrs_dementia_cols = ['f00_nacrs', 'f01_nacrs', 'f02_nacrs', 'f03_nacrs', 'g30_nacrs', 'g310_nacrs', 'g311_nacrs']

    # Claims columns
    claims_cols = ['294_claims', '290_claims', '331_claims']

    # Prescription columns
    rx_cols = ['donepezil prescriptions', 'rivastigmine prescriptions',
               'galantamine prescriptions', 'memantine prescriptions']

    results = []

    for group_name, status_values in cognitive_status_groups.items():
        subset = df[df['cognitive_status'].isin(status_values)]
        n = len(subset)

        if n == 0:
            print(f"  Warning: No participants found for {group_name}")
            continue

        row = {'Group': group_name, 'N': n}

        # Demographics
        # Age
        if 'age' in subset.columns:
            age_valid = subset['age'].dropna()
            row['Age_Mean'] = age_valid.mean() if len(age_valid) > 0 else np.nan
            row['Age_SD'] = age_valid.std() if len(age_valid) > 0 else np.nan

        # Education years
        if 'education_years' in subset.columns:
            edu_valid = subset['education_years'].dropna()
            row['Education_Years_Mean'] = edu_valid.mean() if len(edu_valid) > 0 else np.nan
            row['Education_Years_SD'] = edu_valid.std() if len(edu_valid) > 0 else np.nan

        # Sex (assuming 0=Female, 1=Male based on typical coding)
        if 'sex' in subset.columns:
            sex_valid = subset['sex'].dropna()
            n_male = (sex_valid == 1).sum()
            n_female = (sex_valid == 0).sum()
            row['Male_N'] = n_male
            row['Male_Pct'] = 100 * n_male / len(sex_valid) if len(sex_valid) > 0 else np.nan
            row['Female_N'] = n_female
            row['Female_Pct'] = 100 * n_female / len(sex_valid) if len(sex_valid) > 0 else np.nan

        # Race/Ethnicity
        if 'race' in subset.columns:
            race_valid = subset['race'].dropna()
            # Assuming certain coding for Caucasian
            n_caucasian = ((race_valid == 0) | (race_valid == 'caucasian') | (race_valid == 'white')).sum()
            n_other = len(race_valid) - n_caucasian
            row['Caucasian_N'] = n_caucasian
            row['Caucasian_Pct'] = 100 * n_caucasian / len(race_valid) if len(race_valid) > 0 else np.nan
            row['Other_Ethnicity_N'] = n_other
            row['Other_Ethnicity_Pct'] = 100 * n_other / len(race_valid) if len(race_valid) > 0 else np.nan

        # Cognitive Performance
        if 'mmsetotal' in subset.columns:
            mmse_valid = subset['mmsetotal'].dropna()
            row['MMSE_Mean'] = mmse_valid.mean() if len(mmse_valid) > 0 else np.nan
            row['MMSE_SD'] = mmse_valid.std() if len(mmse_valid) > 0 else np.nan

        if 'moca_total' in subset.columns:
            moca_valid = subset['moca_total'].dropna()
            row['MoCA_Mean'] = moca_valid.mean() if len(moca_valid) > 0 else np.nan
            row['MoCA_SD'] = moca_valid.std() if len(moca_valid) > 0 else np.nan

        # DAD (Discharge Abstract Database) - Hospital records
        dad_available = [c for c in dad_cols + nacrs_cols if c in subset.columns]
        if dad_available:
            # Calculate total DAD records per person
            subset_dad = subset[dad_available].fillna(0)
            total_dad_per_person = subset_dad.sum(axis=1)

            row['DAD_Any_Record_N'] = (total_dad_per_person > 0).sum()
            row['DAD_Any_Record_Pct'] = 100 * row['DAD_Any_Record_N'] / n
            row['DAD_Records_Mean'] = total_dad_per_person.mean()
            row['DAD_Records_SD'] = total_dad_per_person.std()

            # DAD Dementia codes
            dad_dem_available = [c for c in dad_dementia_cols + nacrs_dementia_cols if c in subset.columns]
            if dad_dem_available:
                subset_dad_dem = subset[dad_dem_available].fillna(0)
                total_dad_dem = subset_dad_dem.sum(axis=1)
                row['DAD_Dementia_Code_N'] = (total_dad_dem > 0).sum()
                row['DAD_Dementia_Code_Pct'] = 100 * row['DAD_Dementia_Code_N'] / n
                row['DAD_Dementia_Records_Mean'] = total_dad_dem.mean()
                row['DAD_Dementia_Records_SD'] = total_dad_dem.std()

        # Physician Claims
        claims_available = [c for c in claims_cols if c in subset.columns]
        if claims_available:
            subset_claims = subset[claims_available].fillna(0)
            total_claims = subset_claims.sum(axis=1)

            row['Claims_Any_Record_N'] = (total_claims > 0).sum()
            row['Claims_Any_Record_Pct'] = 100 * row['Claims_Any_Record_N'] / n
            row['Claims_Records_Mean'] = total_claims.mean()
            row['Claims_Records_SD'] = total_claims.std()

            # Claims with dementia codes (all claims columns are dementia-related)
            row['Claims_Dementia_Code_N'] = (total_claims > 0).sum()
            row['Claims_Dementia_Code_Pct'] = 100 * row['Claims_Dementia_Code_N'] / n
            row['Claims_Dementia_Records_Mean'] = total_claims.mean()
            row['Claims_Dementia_Records_SD'] = total_claims.std()

        # Pharmaceutical Information Network (PIN)
        rx_available = [c for c in rx_cols if c in subset.columns]
        if rx_available:
            subset_rx = subset[rx_available].fillna(0)
            total_rx = subset_rx.sum(axis=1)

            # Any medication record (using total prescriptions if available)
            if 'total prescriptions' in subset.columns:
                total_all_rx = subset['total prescriptions'].fillna(0)
                row['PIN_Any_Medication_N'] = (total_all_rx > 0).sum()
                row['PIN_Any_Medication_Pct'] = 100 * row['PIN_Any_Medication_N'] / n
                row['PIN_Medications_Mean'] = total_all_rx.mean()
                row['PIN_Medications_SD'] = total_all_rx.std()

            # Dementia-specific medications
            row['PIN_Dementia_Medication_N'] = (total_rx > 0).sum()
            row['PIN_Dementia_Medication_Pct'] = 100 * row['PIN_Dementia_Medication_N'] / n
            row['PIN_Dementia_Medications_Mean'] = total_rx.mean()
            row['PIN_Dementia_Medications_SD'] = total_rx.std()

        # Lookback years
        lookback_valid = subset['lookback_years'].dropna()
        if len(lookback_valid) > 0:
            row['Lookback_Max_Mean'] = lookback_valid.max()  # Oldest lookback available
            row['Lookback_Mean_Mean'] = lookback_valid.mean()
            row['Lookback_Mean_SD'] = lookback_valid.std()

        results.append(row)
        print(f"  {group_name}: N={n}")

    # Create DataFrame
    results_df = pd.DataFrame(results)

    # Format for publication
    formatted_rows = []

    for _, row in results_df.iterrows():
        formatted = {
            'Measure': '',
            row['Group']: ''
        }
        formatted_rows.append({'Measure': f"(N={int(row['N'])})", row['Group']: ''})

    # Build the formatted table
    measures = [
        ('Age (years), Mean (SD)', 'Age_Mean', 'Age_SD'),
        ('Years of Education, Mean (SD)', 'Education_Years_Mean', 'Education_Years_SD'),
        ('Sex, n (%)', None, None),
        ('  Male', 'Male_N', 'Male_Pct'),
        ('  Female', 'Female_N', 'Female_Pct'),
        ('Ethnic group, n (%)', None, None),
        ('  Caucasian', 'Caucasian_N', 'Caucasian_Pct'),
        ('  Other Ethnicity', 'Other_Ethnicity_N', 'Other_Ethnicity_Pct'),
        ('Cognitive Performance', None, None),
        ('  MMSE Score, mean (SD)', 'MMSE_Mean', 'MMSE_SD'),
        ('  MoCA Score, mean (SD)', 'MoCA_Mean', 'MoCA_SD'),
        ('Discharge Abstract Database', None, None),
        ('  People with >=1 DAD Record, n (%)', 'DAD_Any_Record_N', 'DAD_Any_Record_Pct'),
        ('  Mean Number of DAD Records, mean (SD)', 'DAD_Records_Mean', 'DAD_Records_SD'),
        ('  People with DAD Dementia Code, n (%)', 'DAD_Dementia_Code_N', 'DAD_Dementia_Code_Pct'),
        ('  Mean DAD Dementia Records, mean (SD)', 'DAD_Dementia_Records_Mean', 'DAD_Dementia_Records_SD'),
        ('Physician Claims', None, None),
        ('  People with >=1 Claims Record, n (%)', 'Claims_Any_Record_N', 'Claims_Any_Record_Pct'),
        ('  Mean Number of Claims Records, mean (SD)', 'Claims_Records_Mean', 'Claims_Records_SD'),
        ('  People with Claims Dementia Code, n (%)', 'Claims_Dementia_Code_N', 'Claims_Dementia_Code_Pct'),
        ('  Mean Claims Dementia Records, mean (SD)', 'Claims_Dementia_Records_Mean', 'Claims_Dementia_Records_SD'),
        ('Pharmaceutical Information Network', None, None),
        ('  People with >=1 PIN Medications, n (%)', 'PIN_Any_Medication_N', 'PIN_Any_Medication_Pct'),
        ('  Mean Number of PIN Medications, mean (SD)', 'PIN_Medications_Mean', 'PIN_Medications_SD'),
        ('  People with PIN Dementia Medications, n (%)', 'PIN_Dementia_Medication_N', 'PIN_Dementia_Medication_Pct'),
        ('  Mean PIN Dementia Medications, mean (SD)', 'PIN_Dementia_Medications_Mean', 'PIN_Dementia_Medications_SD'),
        ('Maximum Lookback Years in Admin Data', None, None),
        ('  Oldest Lookback', 'Lookback_Max_Mean', None),
        ('  Mean Lookback, mean (SD)', 'Lookback_Mean_Mean', 'Lookback_Mean_SD'),
    ]

    # Build formatted table
    formatted_data = []
    groups = list(cognitive_status_groups.keys())

    for measure_name, val_col, secondary_col in measures:
        row_data = {'Measure': measure_name}

        for group in groups:
            group_row = results_df[results_df['Group'] == group]
            if group_row.empty:
                row_data[group] = ''
                continue

            group_row = group_row.iloc[0]

            if val_col is None:
                row_data[group] = ''
            elif 'n (%)' in measure_name or ('_N' in val_col and '_Pct' in str(secondary_col)):
                # Format as n (%)
                n_val = group_row.get(val_col, np.nan)
                pct_val = group_row.get(secondary_col, np.nan)
                if pd.notna(n_val) and pd.notna(pct_val):
                    row_data[group] = f"{int(n_val)} ({pct_val:.1f}%)"
                else:
                    row_data[group] = ''
            elif 'Mean (SD)' in measure_name or 'mean (SD)' in measure_name:
                # Format as mean (SD)
                mean_val = group_row.get(val_col, np.nan)
                sd_val = group_row.get(secondary_col, np.nan)
                if pd.notna(mean_val) and pd.notna(sd_val):
                    row_data[group] = f"{mean_val:.1f} ({sd_val:.1f})"
                elif pd.notna(mean_val):
                    row_data[group] = f"{mean_val:.1f}"
                else:
                    row_data[group] = ''
            else:
                # Single value
                val = group_row.get(val_col, np.nan)
                if pd.notna(val):
                    row_data[group] = f"{val:.1f}"
                else:
                    row_data[group] = ''

        formatted_data.append(row_data)

    # Add N row at the top
    n_row = {'Measure': 'N'}
    for group in groups:
        group_row = results_df[results_df['Group'] == group]
        if not group_row.empty:
            n_row[group] = str(int(group_row.iloc[0]['N']))
        else:
            n_row[group] = ''
    formatted_data.insert(0, n_row)

    # Create final DataFrame
    final_df = pd.DataFrame(formatted_data)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    csv_path = os.path.join(output_dir, f"{table_name}.csv")
    final_df.to_csv(csv_path, index=False)
    print(f"  Saved: {csv_path}")

    # Also save raw data
    raw_path = os.path.join(output_dir, f"{table_name}_raw.csv")
    results_df.to_csv(raw_path, index=False)
    print(f"  Saved: {raw_path}")

    return final_df, results_df


# ==============================================================================
# TABLES 3 & 4: PERFORMANCE METRICS
# ==============================================================================

def load_existing_results(results_dir):
    """
    Load existing results from the main results CSV file.
    This file contains all metrics including PR-AUC with confidence intervals.
    """
    # Primary: Use main results file (has PR-AUC and all metrics)
    main_results_path = os.path.join(results_dir, "results_final_eval",
                                      "final_evaluation_metrics_balanced_accuracy_thresh_tuned_PARALLEL.csv")
    if os.path.exists(main_results_path):
        df = pd.read_csv(main_results_path)
        # Rename columns to match expected format
        column_mapping = {
            'AUC (Point [BS Median (CI)])': 'AUC (95% CI)',
            'PR AUC (Point [BS Median (CI)])': 'PR AUC (95% CI)',
            'Sensitivity (Point [BS Median (CI)])': 'Sensitivity (95% CI)',
            'Specificity (Point [BS Median (CI)])': 'Specificity (95% CI)',
            'PPV (Point [BS Median (CI)])': 'PPV (95% CI)',
            'NPV (Point [BS Median (CI)])': 'NPV (95% CI)',
            'F1 Score (Point [BS Median (CI)])': 'F1 (95% CI)',
            'Balanced Acc. (Point [BS Median (CI)])': 'Balanced Acc. (95% CI)',
        }
        df = df.rename(columns=column_mapping)
        return df

    # Fallback: Try publication table (may not have PR-AUC)
    pub_table_path = os.path.join(results_dir, "publication_figures", "tables", "publication_table.csv")
    if os.path.exists(pub_table_path):
        return pd.read_csv(pub_table_path)

    return None


def parse_metric_with_ci(value_str):
    """
    Parse metric string into components.

    Handles multiple formats:
    - '0.913 (0.864-0.950)' - simple format
    - '0.943 [0.944 (0.901-0.976)]' - main results format (point [median (CI)])
    """
    if pd.isna(value_str) or value_str == 'N/A' or value_str == '' or value_str == 'nan':
        return np.nan, np.nan, np.nan

    try:
        value_str = str(value_str).strip()

        # Handle format: "0.943 [0.944 (0.901-0.976)]"
        if '[' in value_str and '(' in value_str:
            # Extract the CI part from brackets
            bracket_part = value_str.split('[')[1].replace(']', '').strip()
            # Now parse "0.944 (0.901-0.976)"
            if '(' in bracket_part:
                parts = bracket_part.split('(')
                point = float(parts[0].strip())  # Use bootstrap median as point
                ci_part = parts[1].replace(')', '').strip()
                if '-' in ci_part:
                    ci_parts = ci_part.split('-')
                    ci_low = float(ci_parts[0].strip())
                    ci_high = float(ci_parts[1].strip())
                    return point, ci_low, ci_high
            return float(bracket_part.split('(')[0].strip()), np.nan, np.nan

        # Handle simple format: "0.913 (0.864-0.950)"
        elif '(' in value_str:
            parts = value_str.split('(')
            point = float(parts[0].strip())
            ci_part = parts[1].replace(')', '').strip()
            if '-' in ci_part:
                ci_parts = ci_part.split('-')
                ci_low = float(ci_parts[0].strip())
                ci_high = float(ci_parts[1].strip())
                return point, ci_low, ci_high
            return point, np.nan, np.nan
        else:
            return float(value_str), np.nan, np.nan
    except:
        return np.nan, np.nan, np.nan


def format_metric_for_table(value_str):
    """
    Format metric string for publication table.

    Converts "0.943 [0.944 (0.901-0.976)]" to "0.944 (0.901-0.976)"
    """
    if pd.isna(value_str) or value_str == 'N/A' or value_str == '' or value_str == 'nan':
        return 'N/A'

    try:
        value_str = str(value_str).strip()

        # Handle format: "0.943 [0.944 (0.901-0.976)]"
        if '[' in value_str and '(' in value_str:
            bracket_part = value_str.split('[')[1].replace(']', '').strip()
            return bracket_part

        # Already in simple format
        return value_str
    except:
        return 'N/A'


def clean_task_name(task_name):
    """
    Clean task name for publication tables by removing prefixes and numbers.

    Examples:
    - '1. Def Normal vs. Def Dementia' -> 'NCD vs. Dementia'
    - '2. Def+Pos Normal vs. Def+Pos Dementia' -> 'NCD vs. Dementia'
    - '5. Multi: Def Normal vs. Def MCI vs. Def Dementia' -> 'NCD vs. MCI vs. Dementia'
    """
    import re

    # Remove task number prefix (e.g., "1. ", "5. Multi: ")
    cleaned = re.sub(r'^\d+\.\s*(Multi:\s*)?', '', task_name)

    # Remove "Def+Pos " and "Def " prefixes
    cleaned = re.sub(r'Def\+Pos\s+', '', cleaned)
    cleaned = re.sub(r'Def\s+', '', cleaned)

    # Replace "Normal" with "NCD" for publication terminology
    cleaned = cleaned.replace('Normal', 'NCD')

    return cleaned.strip()


def parse_metric_point(value_str):
    """Extract point estimate from a metric string."""
    point, _, _ = parse_metric_with_ci(value_str)
    return point


def build_best_models_by_task(results_df, selection_metric_col='AUC (95% CI)'):
    """
    Build task -> best model mapping from results table.

    Parameters:
    -----------
    results_df : DataFrame
        Results table from load_existing_results()
    selection_metric_col : str
        Metric column used to select the best model per task

    Returns:
    --------
    dict: {task_name: {'model', 'selection_score', 'balanced_accuracy'}}
    """
    best_map = {}
    if results_df is None or results_df.empty:
        return best_map

    for task_name, task_df in results_df.groupby('Task'):
        # Exclude benchmark/algorithm rows when selecting ML model
        ml_df = task_df[~task_df['Model'].astype(str).str.contains('jaakkimainen', case=False, na=False)].copy()
        if ml_df.empty:
            continue

        if selection_metric_col not in ml_df.columns:
            continue

        ml_df['__selection_score'] = ml_df[selection_metric_col].apply(parse_metric_point)
        if 'Balanced Acc. (95% CI)' in ml_df.columns:
            ml_df['__balanced_acc'] = ml_df['Balanced Acc. (95% CI)'].apply(parse_metric_point)
        else:
            ml_df['__balanced_acc'] = np.nan

        ml_df = ml_df.dropna(subset=['__selection_score'])
        if ml_df.empty:
            continue

        best_row = ml_df.sort_values('__selection_score', ascending=False).iloc[0]
        best_map[task_name] = {
            'model': best_row['Model'],
            'selection_score': float(best_row['__selection_score']),
            'balanced_accuracy': float(best_row['__balanced_acc']) if pd.notna(best_row['__balanced_acc']) else np.nan,
        }

    return best_map


def build_task_design_summary(output_base_dir):
    """Summarize task sizes, class imbalance, and EPV from generated reports."""
    class_path = os.path.join(output_base_dir, "statistical_comparisons", "class_imbalance_report.csv")
    sample_path = os.path.join(output_base_dir, "epidemiology_analysis", "sample_size_report.csv")

    if not os.path.exists(class_path):
        return []

    class_df = pd.read_csv(class_path)
    sample_df = pd.read_csv(sample_path) if os.path.exists(sample_path) else pd.DataFrame()
    epv_map = {}
    epv_method_map = {}
    if not sample_df.empty and 'Task' in sample_df.columns:
        epv_map = dict(zip(sample_df['Task'], sample_df.get('EPV', np.nan)))
        if 'EPV_Method' in sample_df.columns:
            epv_method_map = dict(zip(sample_df['Task'], sample_df['EPV_Method']))

    summary_lines = []
    task_order = sorted(class_df['Task'].unique(), key=lambda x: int(str(x).split('.', 1)[0]))
    for task_name in task_order:
        task_rows = class_df[class_df['Task'] == task_name].copy()
        if task_rows.empty:
            continue

        train_n = int(task_rows['Train_Count'].sum())
        test_n = int(task_rows['Test_Count'].sum())
        ratio = float(task_rows['Imbalance_Ratio_Train'].iloc[0])
        class_dist = ", ".join(
            f"{row['Class_Name']}={int(row['Train_Count'])}"
            for _, row in task_rows.sort_values('Class_Label').iterrows()
        )
        epv = epv_map.get(task_name, np.nan)
        epv_str = f"{epv:.1f}" if pd.notna(epv) else "N/A"
        epv_method = epv_method_map.get(task_name, "")
        epv_suffix = ""
        if isinstance(epv_method, str) and epv_method:
            if "Multiclass" in epv_method:
                epv_suffix = " (minority-class)"
            elif "Binary" in epv_method:
                epv_suffix = " (positive-class)"
        summary_lines.append(
            f"{task_name}: Train={train_n}, Test={test_n}, "
            f"Class ratio={ratio:.2f}:1, Train classes [{class_dist}], EPV={epv_str}{epv_suffix}"
        )

    return summary_lines


def build_statistical_test_summary(output_base_dir):
    """Summarize significance findings from DeLong and McNemar files."""
    stats_dir = os.path.join(output_base_dir, "statistical_comparisons")
    if not os.path.exists(stats_dir):
        return {}

    summary = {}
    for prefix, label in [("delong_test_", "DeLong"), ("mcnemar_test_", "McNemar")]:
        files = [f for f in os.listdir(stats_dir) if f.startswith(prefix) and f.endswith(".csv")]
        total_tests = 0
        nominal_sig = 0
        fdr_sig = 0
        for filename in files:
            df = pd.read_csv(os.path.join(stats_dir, filename))
            if 'P_Value' in df.columns:
                total_tests += len(df)
                nominal_sig += int((df['P_Value'] < 0.05).sum())
            if 'Significant_FDR' in df.columns:
                fdr_sig += int(df['Significant_FDR'].fillna(False).sum())

        summary[label] = {
            'files': len(files),
            'total_tests': total_tests,
            'nominal_sig': nominal_sig,
            'fdr_sig': fdr_sig,
        }

    return summary


def generate_performance_metrics_table(results_df, task_filter, table_name, output_dir, include_jaakkimainen=False):
    """
    Generate performance metrics table in the required format.

    Parameters:
    -----------
    results_df : DataFrame
        Results from publication_table.csv
    task_filter : list
        List of task names to include
    table_name : str
        Output table name
    output_dir : str
        Output directory
    include_jaakkimainen : bool
        Whether to include Jaakkimainen reference algorithm
    """
    print(f"\n=== Generating {table_name} ===")

    # Filter to relevant tasks
    filtered_df = results_df[results_df['Task'].isin(task_filter)].copy()

    if filtered_df.empty:
        print(f"  Warning: No results found for tasks: {task_filter}")
        return None

    # Metrics to extract - in the order specified by thesis requirements
    metric_columns = {
        'AUC': 'AUC (95% CI)',
        'Balanced Accuracy': 'Balanced Acc. (95% CI)',
        'Sensitivity': 'Sensitivity (95% CI)',
        'Specificity': 'Specificity (95% CI)',
        'PPV': 'PPV (95% CI)',
        'NPV': 'NPV (95% CI)',
        'F1 Score': 'F1 (95% CI)',
        'PR AUC': 'PR AUC (95% CI)',  # Added PR-AUC
    }

    # Build table rows
    rows = []

    for _, row in filtered_df.iterrows():
        # Clean task name for publication (remove prefixes and numbers)
        task = clean_task_name(row['Task'])
        model = row['Model']

        row_data = {
            'Task': task,
            'Model': model
        }

        for metric_name, col_name in metric_columns.items():
            if col_name in row:
                # Format the metric value for publication
                row_data[metric_name] = format_metric_for_table(row[col_name])
            else:
                row_data[metric_name] = 'N/A'

        rows.append(row_data)

    # Add Jaakkimainen literature reference if requested and not already present.
    has_jaakkimainen = any('jaakkimainen' in str(r.get('Model', '')).lower() for r in rows)
    if include_jaakkimainen and not has_jaakkimainen:
        jaak_row = {
            'Task': clean_task_name(JAAKKIMAINEN_METRICS['task']),
            'Model': JAAKKIMAINEN_METRICS['name'],
            'AUC': 'N/A',
            'Balanced Accuracy': f"{JAAKKIMAINEN_METRICS['balanced_accuracy']:.3f}",
            'Sensitivity': f"{JAAKKIMAINEN_METRICS['sensitivity']:.3f}",
            'Specificity': f"{JAAKKIMAINEN_METRICS['specificity']:.3f}",
            'PPV': 'N/A',
            'NPV': 'N/A',
            'F1 Score': 'N/A',
            'PR AUC': 'N/A',
        }
        rows.append(jaak_row)

    # Create DataFrame
    table_df = pd.DataFrame(rows)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    csv_path = os.path.join(output_dir, f"{table_name}.csv")
    table_df.to_csv(csv_path, index=False)
    print(f"  Saved: {csv_path}")

    return table_df


# ==============================================================================
# FOREST PLOTS (FIGURES 1-3)
# ==============================================================================

def generate_thesis_forest_plot(results_df, task_name, figure_name, output_dir,
                                 include_jaakkimainen=False, metric='balanced_accuracy'):
    """
    Generate a publication-ready forest plot for the thesis.

    Parameters:
    -----------
    results_df : DataFrame
        Results from publication_table.csv
    task_name : str
        Task to plot (e.g., '1. Def Normal vs. Def Dementia')
    figure_name : str
        Output figure name
    output_dir : str
        Output directory
    include_jaakkimainen : bool
        Whether to include Jaakkimainen reference
    metric : str
        Metric to plot ('balanced_accuracy', 'auc', etc.)
    """
    print(f"\n=== Generating {figure_name} ===")
    set_publication_style()

    # Filter to task
    task_df = results_df[results_df['Task'] == task_name].copy()

    if task_df.empty:
        print(f"  Warning: No results found for task: {task_name}")
        return None

    # Map metric to column name
    metric_col_map = {
        'balanced_accuracy': 'Balanced Acc. (95% CI)',
        'auc': 'AUC (95% CI)',
        'sensitivity': 'Sensitivity (95% CI)',
        'specificity': 'Specificity (95% CI)',
    }

    col_name = metric_col_map.get(metric, 'Balanced Acc. (95% CI)')

    # Extract data
    models = []
    point_estimates = []
    ci_lowers = []
    ci_uppers = []

    for _, row in task_df.iterrows():
        model_name = row['Model']
        value_str = row.get(col_name, 'N/A')

        point, ci_low, ci_high = parse_metric_with_ci(value_str)

        if not np.isnan(point):
            models.append(model_name)
            point_estimates.append(point)
            ci_lowers.append(ci_low if not np.isnan(ci_low) else point - 0.05)
            ci_uppers.append(ci_high if not np.isnan(ci_high) else point + 0.05)

    # Add Jaakkimainen if requested
    if include_jaakkimainen and metric == 'balanced_accuracy':
        models.append(JAAKKIMAINEN_METRICS['name'])
        point_estimates.append(JAAKKIMAINEN_METRICS['balanced_accuracy'])
        # No CI available for Jaakkimainen
        ci_lowers.append(JAAKKIMAINEN_METRICS['balanced_accuracy'])
        ci_uppers.append(JAAKKIMAINEN_METRICS['balanced_accuracy'])

    if not models:
        print(f"  Warning: No valid data for forest plot")
        return None

    # Sort by point estimate (descending)
    sorted_indices = np.argsort(point_estimates)[::-1]
    models = [models[i] for i in sorted_indices]
    point_estimates = [point_estimates[i] for i in sorted_indices]
    ci_lowers = [ci_lowers[i] for i in sorted_indices]
    ci_uppers = [ci_uppers[i] for i in sorted_indices]

    n_models = len(models)

    # Create figure - professional sizing with extra space for legend below
    fig_height = max(4, n_models * 0.5)
    if include_jaakkimainen:
        fig_height += 0.8  # Extra space for legend at bottom
    fig, ax = plt.subplots(figsize=(10, fig_height))

    # Y positions
    y_positions = np.arange(n_models)

    # Black and white styling - use fill patterns instead of colors
    # Plot each model
    for i, (model, point, lower, upper) in enumerate(zip(models, point_estimates, ci_lowers, ci_uppers)):
        is_reference = 'Jaakkimainen' in model

        # Error bars (CI) - black lines
        if lower != upper:  # Has CI
            ax.plot([lower, upper], [y_positions[i], y_positions[i]],
                    color='black', linewidth=1.5, solid_capstyle='butt')
            # Add caps to error bars
            cap_height = 0.15
            ax.plot([lower, lower], [y_positions[i] - cap_height, y_positions[i] + cap_height],
                    color='black', linewidth=1.5)
            ax.plot([upper, upper], [y_positions[i] - cap_height, y_positions[i] + cap_height],
                    color='black', linewidth=1.5)

        # Point estimate - filled for ML models, hollow for reference
        if is_reference:
            # Hollow square for reference algorithm
            ax.scatter(point, y_positions[i], marker='s', s=100,
                       facecolor='white', edgecolor='black', linewidth=1.5, zorder=5)
        else:
            # Filled circle for ML models
            ax.scatter(point, y_positions[i], marker='o', s=80,
                       facecolor='black', edgecolor='black', linewidth=1, zorder=5)

        # Add value text to the right - using serif font
        if lower != upper:
            ci_text = f"{point:.3f} ({lower:.3f}-{upper:.3f})"
        else:
            ci_text = f"{point:.3f}"
        ax.text(1.02, y_positions[i], ci_text, transform=ax.get_yaxis_transform(),
                va='center', ha='left', fontsize=9, family='serif')

    # Reference line at 0.5 (random classification)
    ax.axvline(x=0.5, color='black', linestyle='--', linewidth=0.8, alpha=0.6)

    # Metric labels
    metric_labels = {
        'auc': 'AUC-ROC',
        'balanced_accuracy': 'Balanced Accuracy',
        'sensitivity': 'Sensitivity',
        'specificity': 'Specificity',
    }

    # Formatting
    ax.set_yticks(y_positions)
    ax.set_yticklabels(models, fontsize=10, family='serif')
    ax.set_xlabel(metric_labels.get(metric, metric.replace('_', ' ').title()), fontsize=11, family='serif')
    ax.set_xlim(max(0.4, min(ci_lowers) - 0.05), 1.02)
    ax.set_ylim(-0.5, n_models - 0.5)

    # Remove top and right spines
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_linewidth(0.8)
    ax.spines['bottom'].set_linewidth(0.8)

    # Title - clean and professional (use clean_task_name function)
    clean_title = clean_task_name(task_name)
    ax.set_title(clean_title, fontsize=12, family='serif', fontweight='bold', pad=10)

    # Column header for CI values
    ax.text(1.02, n_models + 0.2, f'{metric_labels.get(metric, metric)} (95% CI)',
            transform=ax.get_yaxis_transform(), va='bottom', ha='left',
            fontsize=10, family='serif', style='italic')

    # Legend - positioned below the plot to avoid blocking data
    if include_jaakkimainen:
        legend_elements = [
            plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='black',
                      markeredgecolor='black', markersize=8, label='Machine Learning Models'),
            plt.Line2D([0], [0], marker='s', color='w', markerfacecolor='white',
                      markeredgecolor='black', markersize=8, markeredgewidth=1.5,
                      label='Reference Algorithm (Jaakkimainen)')
        ]
        ax.legend(handles=legend_elements, loc='upper center', bbox_to_anchor=(0.5, -0.12),
                  ncol=2, fontsize=9, frameon=True, edgecolor='black', fancybox=False)

    # Subtle grid on x-axis only
    ax.xaxis.grid(True, alpha=0.3, linestyle='-', linewidth=0.5, color='gray')
    ax.set_axisbelow(True)

    plt.tight_layout()

    # Adjust layout to make room for legend below
    if include_jaakkimainen:
        plt.subplots_adjust(bottom=0.15)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    png_path = os.path.join(output_dir, f"{figure_name}.png")
    pdf_path = os.path.join(output_dir, f"{figure_name}.pdf")

    plt.savefig(png_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.savefig(pdf_path, format='pdf', bbox_inches='tight', facecolor='white')
    plt.close()

    print(f"  Saved: {png_path}")
    print(f"  Saved: {pdf_path}")

    # Save data
    data_df = pd.DataFrame({
        'Model': models,
        'Point_Estimate': point_estimates,
        'CI_Lower': ci_lowers,
        'CI_Upper': ci_uppers
    })
    data_path = os.path.join(output_dir, f"{figure_name}_data.csv")
    data_df.to_csv(data_path, index=False)
    print(f"  Saved: {data_path}")

    return fig


# ==============================================================================
# FIGURE 4: AGGREGATED FEATURE IMPORTANCE
# ==============================================================================

def load_feature_importances(stats_dir):
    """Load all feature importance files from statistical comparisons directory."""
    all_importances = []

    if not os.path.exists(stats_dir):
        return None

    for filename in os.listdir(stats_dir):
        if filename.startswith('feature_importance_') and filename.endswith('.csv'):
            filepath = os.path.join(stats_dir, filename)
            try:
                df = pd.read_csv(filepath)
                # Extract task name from filename
                task_name = filename.replace('feature_importance_', '').replace('.csv', '')
                df['Task'] = task_name
                all_importances.append(df)
            except Exception as e:
                print(f"  Warning: Could not load {filename}: {e}")

    if all_importances:
        return pd.concat(all_importances, ignore_index=True)
    return None


def generate_aggregated_feature_importance_plot(stats_dir, output_dir, top_n=10):
    """
    Generate Figure 4: Aggregated feature importance across all tasks.

    Uses model-based feature importance aggregated across tasks.
    """
    print(f"\n=== Generating Figure 4: Aggregated Feature Importance ===")
    set_publication_style()

    # Load feature importances
    importance_df = load_feature_importances(stats_dir)

    if importance_df is None or importance_df.empty:
        print("  Warning: No feature importance data found")
        # Create placeholder with message
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.text(0.5, 0.5, 'Feature importance data not available.\nRe-run analysis to generate.',
                ha='center', va='center', fontsize=14, transform=ax.transAxes)
        ax.set_axis_off()

        os.makedirs(output_dir, exist_ok=True)
        plt.savefig(os.path.join(output_dir, "Figure_4_Feature_Importance.png"),
                    dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        return None

    # Aggregate importance across tasks and models
    # Use normalized importance and average across all occurrences
    agg_importance = importance_df.groupby('Feature').agg({
        'Importance': 'mean',
        'Importance_Normalized': 'mean'
    }).reset_index()

    # Sort and get top N
    agg_importance = agg_importance.sort_values('Importance_Normalized', ascending=False).head(top_n)

    # Create figure - professional sizing
    fig, ax = plt.subplots(figsize=(10, max(5, top_n * 0.45)))

    # Horizontal bar plot - black and white with gradient shading
    y_pos = np.arange(len(agg_importance))

    # Use grayscale gradient based on importance
    importance_vals = agg_importance['Importance_Normalized'].values
    max_imp = max(importance_vals)
    gray_values = [str(0.3 + 0.5 * (1 - val/max_imp)) for val in importance_vals]

    bars = ax.barh(y_pos, importance_vals,
                   color=gray_values, edgecolor='black', linewidth=0.8)

    # Clean up feature names for display - more professional
    feature_names = agg_importance['Feature'].values
    clean_names = []
    feature_name_map = {
        '331_claims': 'ICD-9 331 (Physician Claims)',
        '290_claims': 'ICD-9 290 (Physician Claims)',
        '294_claims': 'ICD-9 294 (Physician Claims)',
        'g30_nacrs': 'ICD-10 G30 (NACRS)',
        'g310_nacrs': 'ICD-10 G31.0 (NACRS)',
        'g31_sum_nacrs': 'ICD-10 G31 Sum (NACRS)',
        'f0_sum_nacrs': 'ICD-10 F0 Sum (NACRS)',
        'f03_nacrs': 'ICD-10 F03 (NACRS)',
        'total prescriptions': 'Total Prescriptions',
        'donepezil prescriptions': 'Donepezil Prescriptions',
    }
    for name in feature_names:
        clean_name = feature_name_map.get(str(name), str(name).replace('_', ' ').title())
        if len(clean_name) > 35:
            clean_name = clean_name[:32] + '...'
        clean_names.append(clean_name)

    # Labels - professional serif font
    ax.set_yticks(y_pos)
    ax.set_yticklabels(clean_names, fontsize=10, family='serif')
    ax.set_xlabel('Mean Normalized Importance', fontsize=11, family='serif')
    ax.set_title('Top 10 Most Important Predictive Features',
                 fontsize=12, family='serif', fontweight='bold', pad=10)

    # Add value labels on bars
    for i, (bar, val) in enumerate(zip(bars, importance_vals)):
        ax.text(val + 0.003, bar.get_y() + bar.get_height()/2,
                f'{val:.3f}', va='center', ha='left', fontsize=9, family='serif')

    ax.set_xlim(0, max(importance_vals) * 1.15)
    ax.invert_yaxis()  # Top feature at top

    # Remove top and right spines
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_linewidth(0.8)
    ax.spines['bottom'].set_linewidth(0.8)

    # Subtle grid on x-axis only
    ax.xaxis.grid(True, alpha=0.3, linestyle='-', linewidth=0.5, color='gray')
    ax.set_axisbelow(True)

    # Add explanatory note below the figure
    note_text = ("Note: Feature importance values represent mean normalized model-based importance\n"
                 "aggregated across all six classification tasks and all models (tree gain/Gini importance\n"
                 "or coefficient magnitude, depending on model family).")
    fig.text(0.5, -0.02, note_text, ha='center', va='top', fontsize=9, family='serif',
             style='italic', wrap=True)

    plt.tight_layout()
    plt.subplots_adjust(bottom=0.18)  # Make room for note

    # Save
    os.makedirs(output_dir, exist_ok=True)
    png_path = os.path.join(output_dir, "Figure_7_Feature_Importance.png")
    pdf_path = os.path.join(output_dir, "Figure_7_Feature_Importance.pdf")

    plt.savefig(png_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.savefig(pdf_path, format='pdf', bbox_inches='tight', facecolor='white')
    plt.close()

    print(f"  Saved: {png_path}")
    print(f"  Saved: {pdf_path}")

    # Save data
    data_path = os.path.join(output_dir, "Figure_7_Feature_Importance_data.csv")
    agg_importance.to_csv(data_path, index=False)
    print(f"  Saved: {data_path}")

    return fig


# ==============================================================================
# CONFUSION MATRICES
# ==============================================================================

def generate_confusion_matrix_figure(confusion_matrix_dir, output_dir, task_num, task_name,
                                     model_name, figure_num, is_definite=True,
                                     selection_metric_label='AUC-ROC',
                                     selection_score=np.nan,
                                     balanced_acc=np.nan):
    """
    Generate a normalized confusion matrix figure for the best performing model.

    Parameters:
    -----------
    confusion_matrix_dir : str
        Directory containing confusion matrix CSV files
    output_dir : str
        Output directory for figures
    task_num : int
        Task number (1-6)
    task_name : str
        Full task name
    model_name : str
        Name of the best performing model
    figure_num : int
        Figure number for output filename
    is_definite : bool
        Whether this is a definite (True) or def+pos (False) task
    """
    print(f"\n=== Generating Figure {figure_num}: Confusion Matrix for Task {task_num} ===")
    set_publication_style()

    # Find matching file
    normalized_file = None
    for filename in os.listdir(confusion_matrix_dir):
        if filename.startswith(f'confusion_matrix_{task_num}_') and '_normalized.csv' in filename:
            # Check if it matches the model name (approximately)
            model_check = model_name.replace(' ', '_').replace('(', '').replace(')', '')
            if model_check.lower() in filename.lower() or model_name.split()[0].lower() in filename.lower():
                normalized_file = os.path.join(confusion_matrix_dir, filename)
                break

    if normalized_file is None or not os.path.exists(normalized_file):
        print(f"  Warning: No matching confusion matrix file for Task {task_num} and model '{model_name}'")
        return None

    # Load confusion matrix
    try:
        cm_df = pd.read_csv(normalized_file, index_col=0)
    except Exception as e:
        print(f"  Error loading confusion matrix: {e}")
        return None

    # Create figure
    cm_array = cm_df.values
    class_labels = cm_df.columns.tolist()

    # Determine what "Impaired" means based on task
    # For NCD vs Dementia tasks (1, 2): Impaired = Dementia
    # For NCD vs MCI tasks (3, 4): Impaired = MCI
    if task_num in [1, 2]:
        impaired_label = 'Dementia'
    elif task_num in [3, 4]:
        impaired_label = 'MCI'
    else:
        impaired_label = 'Impaired'  # Multiclass won't have this

    # Clean up class labels
    label_map = {
        'Normal': 'NCD',
        'Impaired': impaired_label,
        'Dementia': 'Dementia',
        'MCI': 'MCI',
    }
    clean_labels = [label_map.get(l, l) for l in class_labels]

    fig, ax = plt.subplots(figsize=(6, 5))

    # Plot heatmap in grayscale
    im = ax.imshow(cm_array, cmap='Greys', aspect='auto', vmin=0, vmax=1)

    # Add colorbar
    cbar = ax.figure.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    cbar.ax.set_ylabel('Proportion', rotation=-90, va="bottom", fontsize=10, family='serif')
    cbar.ax.tick_params(labelsize=9)

    # Set ticks and labels
    ax.set_xticks(np.arange(len(clean_labels)))
    ax.set_yticks(np.arange(len(clean_labels)))
    ax.set_xticklabels(clean_labels, fontsize=10, family='serif')
    ax.set_yticklabels(clean_labels, fontsize=10, family='serif')

    # Rotate x labels
    plt.setp(ax.get_xticklabels(), rotation=45, ha="right", rotation_mode="anchor")

    # Add text annotations
    thresh = 0.5
    for i in range(len(clean_labels)):
        for j in range(len(clean_labels)):
            val = cm_array[i, j]
            text_color = "white" if val > thresh else "black"
            ax.text(j, i, f'{val:.2f}', ha="center", va="center",
                    color=text_color, fontsize=11, family='serif', fontweight='bold')

    # Labels and title
    ax.set_xlabel('Predicted Label', fontsize=11, family='serif')
    ax.set_ylabel('True Label', fontsize=11, family='serif')

    # Clean title
    clean_task = clean_task_name(task_name)
    metric_parts = []
    if pd.notna(selection_score):
        metric_parts.append(f"{selection_metric_label} = {selection_score:.3f}")
    if pd.notna(balanced_acc):
        metric_parts.append(f"Balanced Acc. = {balanced_acc:.3f}")
    metric_text = ", ".join(metric_parts)
    title_suffix = f" ({model_name}{', ' + metric_text if metric_text else ''})"
    ax.set_title(f'{clean_task}\n{title_suffix}',
                 fontsize=11, family='serif', fontweight='bold', pad=10)

    plt.tight_layout()

    # Save
    os.makedirs(output_dir, exist_ok=True)
    prefix = "Definite" if is_definite else "DefPos"
    png_path = os.path.join(output_dir, f"Figure_{figure_num}_{prefix}_Confusion_Matrix_Task{task_num}.png")
    pdf_path = os.path.join(output_dir, f"Figure_{figure_num}_{prefix}_Confusion_Matrix_Task{task_num}.pdf")
    csv_path = os.path.join(output_dir, f"Figure_{figure_num}_{prefix}_Confusion_Matrix_Task{task_num}_data.csv")

    plt.savefig(png_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.savefig(pdf_path, format='pdf', bbox_inches='tight', facecolor='white')
    plt.close()

    # Save corresponding data table with cleaned labels
    cm_df_clean = cm_df.copy()
    cm_df_clean.columns = clean_labels
    cm_df_clean.index = clean_labels
    cm_df_clean.to_csv(csv_path)

    print(f"  Saved: {png_path}")
    print(f"  Saved: {pdf_path}")
    print(f"  Saved: {csv_path}")

    return fig


def generate_all_confusion_matrices(confusion_matrix_dir, output_dir,
                                    best_models_by_task,
                                    selection_metric_label='AUC-ROC'):
    """Generate all 6 confusion matrices for best performing models.

    Returns:
    --------
    dict : Dictionary of generated file paths
    """
    print("\n" + "="*70)
    print("GENERATING CONFUSION MATRICES FOR BEST PERFORMING MODELS")
    print("="*70)

    generated_files = {}

    # Figure numbers: 8-10 for Definite, 11-13 for Def+Pos
    figure_configs = [
        # (task_num, task_name, figure_num, is_definite)
        (1, '1. Def Normal vs. Def Dementia', 8, True),
        (3, '3. Def Normal vs. Def MCI', 9, True),
        (5, '5. Multi: Def Normal vs. Def MCI vs. Def Dementia', 10, True),
        (2, '2. Def+Pos Normal vs. Def+Pos Dementia', 11, False),
        (4, '4. Def+Pos Normal vs. Def+Pos MCI', 12, False),
        (6, '6. Multi: Def+Pos Normal vs. Def+Pos MCI vs. Def+Pos Dementia', 13, False),
    ]

    for task_num, task_name, figure_num, is_definite in figure_configs:
        best_info = best_models_by_task.get(task_name, {})
        model_name = best_info.get('model')
        selection_score = best_info.get('selection_score', np.nan)
        balanced_acc = best_info.get('balanced_accuracy', np.nan)

        if not model_name:
            print(f"  WARNING: No best-model mapping found for {task_name}")
            continue

        prefix = "Definite" if is_definite else "DefPos"
        try:
            fig = generate_confusion_matrix_figure(
                confusion_matrix_dir, output_dir, task_num, task_name,
                model_name, figure_num, is_definite,
                selection_metric_label=selection_metric_label,
                selection_score=selection_score,
                balanced_acc=balanced_acc
            )
            if fig is None:
                continue
            # Track generated files
            generated_files[f'figure_{figure_num}_cm'] = os.path.join(
                output_dir, f"Figure_{figure_num}_{prefix}_Confusion_Matrix_Task{task_num}.png"
            )
            generated_files[f'figure_{figure_num}_cm_data'] = os.path.join(
                output_dir, f"Figure_{figure_num}_{prefix}_Confusion_Matrix_Task{task_num}_data.csv"
            )
        except Exception as e:
            print(f"  WARNING: Figure {figure_num} confusion matrix failed: {e}")

    return generated_files


# ==============================================================================
# WORD DOCUMENT GENERATION (VANCOUVER STYLE)
# ==============================================================================

def set_cell_shading(cell, fill_color):
    """Set background shading for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table):
    """Set borders for entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Check if borders already exist
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)


def generate_thesis_word_document(output_dir, thesis_output_dir=None,
                                  best_models_by_task=None,
                                  selection_metric_label='AUC-ROC'):
    """
    Generate a unified Word document with all thesis tables and figures.
    Formatted in Vancouver style for publication.

    Parameters:
    -----------
    output_dir : str
        Base output directory containing results
    thesis_output_dir : str, optional
        Directory containing thesis outputs (defaults to output_dir/thesis_outputs)

    Returns:
    --------
    str : Path to generated Word document
    """
    if not DOCX_AVAILABLE:
        print("ERROR: python-docx not available. Cannot generate Word document.")
        return None

    print("\n" + "=" * 70)
    print("GENERATING UNIFIED WORD DOCUMENT (VANCOUVER STYLE)")
    print("=" * 70)

    if thesis_output_dir is None:
        thesis_output_dir = os.path.join(output_dir, "thesis_outputs")

    if best_models_by_task is None:
        best_models_by_task = {}

    # Create document
    doc = Document()

    # ===========================================================================
    # DOCUMENT STYLES SETUP
    # ===========================================================================

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title style
    title_style = doc.styles.add_style('ThesisTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(24)

    # Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(12)

    # Table/Figure caption style
    caption_style = doc.styles.add_style('TableCaption', WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.name = 'Times New Roman'
    caption_style.font.size = Pt(11)
    caption_style.font.bold = True
    caption_style.paragraph_format.space_before = Pt(12)
    caption_style.paragraph_format.space_after = Pt(6)

    # Table note style
    note_style = doc.styles.add_style('TableNote', WD_STYLE_TYPE.PARAGRAPH)
    note_style.font.name = 'Times New Roman'
    note_style.font.size = Pt(10)
    note_style.font.italic = True
    note_style.paragraph_format.space_before = Pt(6)
    note_style.paragraph_format.space_after = Pt(12)

    def format_cm_caption(task_name, base_text):
        """Attach selected-model metadata to confusion-matrix caption."""
        info = best_models_by_task.get(task_name, {})
        model_name = info.get('model')
        selection_score = info.get('selection_score', np.nan)
        balanced_acc = info.get('balanced_accuracy', np.nan)

        if not model_name:
            return base_text

        metric_parts = []
        if pd.notna(selection_score):
            metric_parts.append(f"{selection_metric_label} = {selection_score:.3f}")
        if pd.notna(balanced_acc):
            metric_parts.append(f"Balanced Acc. = {balanced_acc:.3f}")
        metrics_text = ", ".join(metric_parts)
        return f"{base_text} ({model_name}{', ' + metrics_text if metrics_text else ''})"

    # ===========================================================================
    # TITLE PAGE
    # ===========================================================================

    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph("Machine Learning Classification of Cognitive Disorders\nUsing Administrative Health Data", style='ThesisTitle')
    doc.add_paragraph()

    subtitle = doc.add_paragraph("Tables and Figures for Thesis")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)

    doc.add_paragraph()

    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ===========================================================================
    # TABLE OF CONTENTS
    # ===========================================================================

    doc.add_heading('Table of Contents', level=1)

    toc_items = [
        ("SECTION A: DEFINITE COGNITIVE STATUS ANALYSES", ""),
        ("  Table 1: Participant Characteristics (Definite)", ""),
        ("  Table 2: NCD vs Dementia Classification (Definite)", ""),
        ("  Table 3: NCD vs MCI Classification (Definite)", ""),
        ("  Table 4: Multiclass Classification (Definite)", ""),
        ("  Figure 1: Forest Plot – Definite NCD vs Dementia", ""),
        ("  Figure 2: Forest Plot – Definite NCD vs MCI", ""),
        ("  Figure 3: Forest Plot – Definite Multiclass", ""),
        ("", ""),
        ("SECTION B: DEFINITE + POSSIBLE COGNITIVE STATUS ANALYSES", ""),
        ("  Table 5: Participant Characteristics (Definite + Possible)", ""),
        ("  Table 6: NCD vs Dementia Classification (Definite + Possible)", ""),
        ("  Table 7: NCD vs MCI Classification (Definite + Possible)", ""),
        ("  Table 8: Multiclass Classification (Definite + Possible)", ""),
        ("  Figure 4: Forest Plot – Def+Pos NCD vs Dementia", ""),
        ("  Figure 5: Forest Plot – Def+Pos NCD vs MCI", ""),
        ("  Figure 6: Forest Plot – Def+Pos Multiclass", ""),
        ("", ""),
        ("SECTION C: SUPPLEMENTARY ANALYSES", ""),
        ("  Figure 7: Aggregated Feature Importance", ""),
    ]

    for item, _ in toc_items:
        if item:
            toc_para = doc.add_paragraph()
            if item.startswith("SECTION"):
                run = toc_para.add_run(item)
                run.bold = True
            else:
                toc_para.add_run(item)

    doc.add_page_break()

    # ===========================================================================
    # SECTION A: DEFINITE COGNITIVE STATUS ANALYSES
    # ===========================================================================

    section_a = doc.add_heading('SECTION A: Definite Cognitive Status Analyses', level=1)

    # ---------------------------------------------------------------------------
    # TABLE 1: PARTICIPANT CHARACTERISTICS (DEFINITE)
    # ---------------------------------------------------------------------------

    table1_path = os.path.join(thesis_output_dir, "Table_1_Participant_Characteristics_Definite.csv")

    if os.path.exists(table1_path):
        table1_df = pd.read_csv(table1_path)

        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Table 1. ").bold = True
        caption.add_run("Participant characteristics by definite cognitive status")

        table = doc.add_table(rows=1, cols=len(table1_df.columns))
        table.style = 'Table Grid'
        set_table_borders(table)

        header_cells = table.rows[0].cells
        for i, col_name in enumerate(table1_df.columns):
            header_cells[i].text = col_name
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            set_cell_shading(header_cells[i], 'D9D9D9')

        for _, row in table1_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ''
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
                row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

        note = doc.add_paragraph(style='TableNote')
        note.add_run("Abbreviations: ").bold = True
        note.add_run("NCD, no cognitive disorder; MCI, mild cognitive impairment; DAD, Discharge Abstract Database; PIN, Pharmaceutical Information Network; MMSE, Mini-Mental State Examination; MoCA, Montreal Cognitive Assessment. Values are presented as mean (SD) for continuous variables and n (%) for categorical variables.")
    else:
        doc.add_paragraph("Table 1 data not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # TABLE 2: NCD vs DEMENTIA (DEFINITE ONLY)
    # ---------------------------------------------------------------------------

    table2_path = os.path.join(thesis_output_dir, "Table_2_Definite_NCD_vs_Dementia.csv")

    if os.path.exists(table2_path):
        table2_df = pd.read_csv(table2_path)

        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Table 2. ").bold = True
        caption.add_run("NCD versus dementia classification performance for definite cognitive status with 95% confidence intervals")

        table = doc.add_table(rows=1, cols=len(table2_df.columns))
        table.style = 'Table Grid'
        set_table_borders(table)

        header_cells = table.rows[0].cells
        for i, col_name in enumerate(table2_df.columns):
            header_cells[i].text = col_name
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
            header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            set_cell_shading(header_cells[i], 'D9D9D9')

        for _, row in table2_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ''
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
                row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
                if 'Jaakkimainen' in str(row.get('Model', '')):
                    set_cell_shading(row_cells[i], 'E8E8E8')

        note = doc.add_paragraph(style='TableNote')
        note.add_run("Abbreviations: ").bold = True
        note.add_run("AUC, area under the receiver operating characteristic curve; PPV, positive predictive value; NPV, negative predictive value; PR AUC, precision-recall AUC. The Jaakkimainen Algorithm (shaded) represents the reference administrative data case definition. Values in parentheses are 95% bootstrap confidence intervals.")
    else:
        doc.add_paragraph("Table 2 data not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # TABLE 3: NCD vs MCI (DEFINITE ONLY)
    # ---------------------------------------------------------------------------

    table3_path = os.path.join(thesis_output_dir, "Table_3_Definite_NCD_vs_MCI.csv")

    if os.path.exists(table3_path):
        table3_df = pd.read_csv(table3_path)

        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Table 3. ").bold = True
        caption.add_run("NCD versus MCI classification performance for definite cognitive status with 95% confidence intervals")

        table = doc.add_table(rows=1, cols=len(table3_df.columns))
        table.style = 'Table Grid'
        set_table_borders(table)

        header_cells = table.rows[0].cells
        for i, col_name in enumerate(table3_df.columns):
            header_cells[i].text = col_name
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
            header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            set_cell_shading(header_cells[i], 'D9D9D9')

        for _, row in table3_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ''
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
                row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

        note = doc.add_paragraph(style='TableNote')
        note.add_run("Abbreviations: ").bold = True
        note.add_run("AUC, area under the receiver operating characteristic curve; PPV, positive predictive value; NPV, negative predictive value; PR AUC, precision-recall AUC. Values in parentheses are 95% bootstrap confidence intervals.")
    else:
        doc.add_paragraph("Table 3 data not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # TABLE 4: MULTICLASS CLASSIFICATION (DEFINITE ONLY)
    # ---------------------------------------------------------------------------

    table4_path = os.path.join(thesis_output_dir, "Table_4_Definite_Multiclass.csv")

    if os.path.exists(table4_path):
        table4_df = pd.read_csv(table4_path)

        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Table 4. ").bold = True
        caption.add_run("Multiclass classification performance for definite cognitive status (NCD vs MCI vs Dementia) with 95% confidence intervals")

        table = doc.add_table(rows=1, cols=len(table4_df.columns))
        table.style = 'Table Grid'
        set_table_borders(table)

        header_cells = table.rows[0].cells
        for i, col_name in enumerate(table4_df.columns):
            header_cells[i].text = col_name
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
            header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            set_cell_shading(header_cells[i], 'D9D9D9')

        for _, row in table4_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ''
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
                row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

        note = doc.add_paragraph(style='TableNote')
        note.add_run("Note: ").bold = True
        note.add_run("Multiclass metrics use macro-averaging across NCD, MCI, and Dementia classes. N/A indicates model convergence issues. Values in parentheses are 95% bootstrap confidence intervals.")
    else:
        doc.add_paragraph("Table 4 data not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # FIGURE 1: FOREST PLOT - DEFINITE NCD vs DEMENTIA
    # ---------------------------------------------------------------------------

    fig1_path = os.path.join(thesis_output_dir, "Figure_1_Definite_NCD_vs_Dementia.png")

    if os.path.exists(fig1_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 1. ").bold = True
        caption.add_run("Forest plot: Balanced accuracy for definite NCD versus definite dementia classification")

        doc.add_picture(fig1_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        legend = doc.add_paragraph(style='TableNote')
        legend.add_run("Note: ").bold = True
        legend.add_run("Filled circles represent machine learning models; hollow square represents the Jaakkimainen reference algorithm. Horizontal lines indicate 95% bootstrap confidence intervals. Dashed vertical line indicates random classification (0.5).")
    else:
        doc.add_paragraph("Figure 1 not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # FIGURE 2: FOREST PLOT - DEFINITE NCD vs MCI
    # ---------------------------------------------------------------------------

    fig2_path = os.path.join(thesis_output_dir, "Figure_2_Definite_NCD_vs_MCI.png")

    if os.path.exists(fig2_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 2. ").bold = True
        caption.add_run("Forest plot: Balanced accuracy for definite NCD versus definite MCI classification")

        doc.add_picture(fig2_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        legend = doc.add_paragraph(style='TableNote')
        legend.add_run("Note: ").bold = True
        legend.add_run("Filled circles represent model point estimates with 95% bootstrap confidence intervals. NCD versus MCI classification is more challenging than NCD versus dementia, reflected in lower balanced accuracy values.")
    else:
        doc.add_paragraph("Figure 2 not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # FIGURE 3: FOREST PLOT - DEFINITE MULTICLASS
    # ---------------------------------------------------------------------------

    fig3_path = os.path.join(thesis_output_dir, "Figure_3_Definite_Multiclass.png")

    if os.path.exists(fig3_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 3. ").bold = True
        caption.add_run("Forest plot: Balanced accuracy for definite three-way classification (NCD vs MCI vs Dementia)")

        doc.add_picture(fig3_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        legend = doc.add_paragraph(style='TableNote')
        legend.add_run("Note: ").bold = True
        legend.add_run("Balanced accuracy represents macro-averaged sensitivity across all three classes. Ensemble methods demonstrated the most consistent performance.")
    else:
        doc.add_paragraph("Figure 3 not available.")

    doc.add_page_break()

    # ===========================================================================
    # SECTION B: DEFINITE + POSSIBLE COGNITIVE STATUS ANALYSES
    # ===========================================================================

    doc.add_heading('SECTION B: Definite + Possible Cognitive Status Analyses', level=1)

    # ---------------------------------------------------------------------------
    # TABLE 5: PARTICIPANT CHARACTERISTICS (DEFINITE + POSSIBLE)
    # ---------------------------------------------------------------------------

    table5_char_path = os.path.join(thesis_output_dir, "Table_1b_Participant_Characteristics_DefPos.csv")

    if os.path.exists(table5_char_path):
        table5_char_df = pd.read_csv(table5_char_path)

        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Table 5. ").bold = True
        caption.add_run("Participant characteristics by definite and possible cognitive status combined")

        table = doc.add_table(rows=1, cols=len(table5_char_df.columns))
        table.style = 'Table Grid'
        set_table_borders(table)

        header_cells = table.rows[0].cells
        for i, col_name in enumerate(table5_char_df.columns):
            header_cells[i].text = col_name
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            set_cell_shading(header_cells[i], 'D9D9D9')

        for _, row in table5_char_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ''
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
                row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

        note = doc.add_paragraph(style='TableNote')
        note.add_run("Note: ").bold = True
        note.add_run("This table combines participants with definite and possible classifications. Sample sizes are larger than Table 1 due to inclusion of possible cases.")
    else:
        doc.add_paragraph("Table 5 data not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # TABLE 6: NCD vs DEMENTIA (DEFINITE + POSSIBLE)
    # ---------------------------------------------------------------------------

    table6_dem_path = os.path.join(thesis_output_dir, "Table_6_DefPos_NCD_vs_Dementia.csv")

    if os.path.exists(table6_dem_path):
        table6_dem_df = pd.read_csv(table6_dem_path)

        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Table 6. ").bold = True
        caption.add_run("NCD versus dementia classification performance for definite + possible cognitive status with 95% confidence intervals")

        table = doc.add_table(rows=1, cols=len(table6_dem_df.columns))
        table.style = 'Table Grid'
        set_table_borders(table)

        header_cells = table.rows[0].cells
        for i, col_name in enumerate(table6_dem_df.columns):
            header_cells[i].text = col_name
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
            header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            set_cell_shading(header_cells[i], 'D9D9D9')

        for _, row in table6_dem_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ''
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
                row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

        note = doc.add_paragraph(style='TableNote')
        note.add_run("Abbreviations: ").bold = True
        note.add_run("AUC, area under the receiver operating characteristic curve; PPV, positive predictive value; NPV, negative predictive value; PR AUC, precision-recall AUC. Values in parentheses are 95% bootstrap confidence intervals.")
    else:
        doc.add_paragraph("Table 6 data not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # TABLE 7: NCD vs MCI (DEFINITE + POSSIBLE)
    # ---------------------------------------------------------------------------

    table7_mci_path = os.path.join(thesis_output_dir, "Table_7_DefPos_NCD_vs_MCI.csv")

    if os.path.exists(table7_mci_path):
        table7_mci_df = pd.read_csv(table7_mci_path)

        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Table 7. ").bold = True
        caption.add_run("NCD versus MCI classification performance for definite + possible cognitive status with 95% confidence intervals")

        table = doc.add_table(rows=1, cols=len(table7_mci_df.columns))
        table.style = 'Table Grid'
        set_table_borders(table)

        header_cells = table.rows[0].cells
        for i, col_name in enumerate(table7_mci_df.columns):
            header_cells[i].text = col_name
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
            header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            set_cell_shading(header_cells[i], 'D9D9D9')

        for _, row in table7_mci_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ''
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
                row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

        note = doc.add_paragraph(style='TableNote')
        note.add_run("Abbreviations: ").bold = True
        note.add_run("AUC, area under the receiver operating characteristic curve; PPV, positive predictive value; NPV, negative predictive value; PR AUC, precision-recall AUC. Values in parentheses are 95% bootstrap confidence intervals.")
    else:
        doc.add_paragraph("Table 7 data not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # TABLE 8: MULTICLASS CLASSIFICATION (DEFINITE + POSSIBLE)
    # ---------------------------------------------------------------------------

    table8_path = os.path.join(thesis_output_dir, "Table_8_DefPos_Multiclass.csv")

    if os.path.exists(table8_path):
        table8_df = pd.read_csv(table8_path)

        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Table 8. ").bold = True
        caption.add_run("Multiclass classification performance for definite + possible cognitive status (NCD vs MCI vs Dementia) with 95% confidence intervals")

        table = doc.add_table(rows=1, cols=len(table8_df.columns))
        table.style = 'Table Grid'
        set_table_borders(table)

        header_cells = table.rows[0].cells
        for i, col_name in enumerate(table8_df.columns):
            header_cells[i].text = col_name
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
            header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            set_cell_shading(header_cells[i], 'D9D9D9')

        for _, row in table8_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ''
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
                row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

        note = doc.add_paragraph(style='TableNote')
        note.add_run("Note: ").bold = True
        note.add_run("Multiclass metrics for expanded cohort. N/A indicates model convergence issues. Values in parentheses are 95% bootstrap confidence intervals.")
    else:
        doc.add_paragraph("Table 8 data not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # FIGURE 4: FOREST PLOT - DEF+POS NCD vs DEMENTIA
    # ---------------------------------------------------------------------------

    fig4_path = os.path.join(thesis_output_dir, "Figure_4_DefPos_NCD_vs_Dementia.png")

    if os.path.exists(fig4_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 4. ").bold = True
        caption.add_run("Forest plot: Balanced accuracy for definite + possible NCD versus dementia classification")

        doc.add_picture(fig4_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        legend = doc.add_paragraph(style='TableNote')
        legend.add_run("Note: ").bold = True
        legend.add_run("Filled circles represent model point estimates with 95% bootstrap confidence intervals. Results for expanded cohort including both definite and possible cognitive status classifications.")
    else:
        doc.add_paragraph("Figure 4 not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # FIGURE 5: FOREST PLOT - DEF+POS NCD vs MCI
    # ---------------------------------------------------------------------------

    fig5_path = os.path.join(thesis_output_dir, "Figure_5_DefPos_NCD_vs_MCI.png")

    if os.path.exists(fig5_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 5. ").bold = True
        caption.add_run("Forest plot: Balanced accuracy for definite + possible NCD versus MCI classification")

        doc.add_picture(fig5_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        legend = doc.add_paragraph(style='TableNote')
        legend.add_run("Note: ").bold = True
        legend.add_run("NCD versus MCI classification remains challenging even with expanded sample sizes. Filled circles represent model point estimates with 95% bootstrap confidence intervals.")
    else:
        doc.add_paragraph("Figure 5 not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # FIGURE 6: FOREST PLOT - DEF+POS MULTICLASS
    # ---------------------------------------------------------------------------

    fig6_path = os.path.join(thesis_output_dir, "Figure_6_DefPos_Multiclass.png")

    if os.path.exists(fig6_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 6. ").bold = True
        caption.add_run("Forest plot: Balanced accuracy for definite + possible three-way classification (NCD vs MCI vs Dementia)")

        doc.add_picture(fig6_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        legend = doc.add_paragraph(style='TableNote')
        legend.add_run("Note: ").bold = True
        legend.add_run("Balanced accuracy represents macro-averaged sensitivity across all three classes for the expanded cohort.")
    else:
        doc.add_paragraph("Figure 6 not available.")

    doc.add_page_break()

    # ===========================================================================
    # SECTION C: SUPPLEMENTARY ANALYSES
    # ===========================================================================

    doc.add_heading('SECTION C: Supplementary Analyses', level=1)

    # ---------------------------------------------------------------------------
    # FIGURE 7: FEATURE IMPORTANCE
    # ---------------------------------------------------------------------------

    fig7_path = os.path.join(thesis_output_dir, "Figure_7_Feature_Importance.png")

    if os.path.exists(fig7_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 7. ").bold = True
        caption.add_run("Top 10 most important predictive features aggregated across all classification tasks")

        doc.add_picture(fig7_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        legend = doc.add_paragraph(style='TableNote')
        legend.add_run("Note: ").bold = True
        legend.add_run("Feature importance values represent mean normalized model-based importance aggregated across all six classification tasks and all models (tree gain/Gini importance or coefficient magnitude, depending on model family). ICD-9 codes (331, 290, 294) from physician claims and ICD-10 codes (G30, F0 series) from NACRS emerged as strongest predictors. Donepezil prescriptions (cholinesterase inhibitor) also demonstrated predictive value.")
    else:
        doc.add_paragraph("Figure 7 not available.")

    doc.add_page_break()

    # ===========================================================================
    # SECTION D: CONFUSION MATRICES FOR BEST PERFORMING MODELS
    # ===========================================================================

    doc.add_heading('SECTION D: Confusion Matrices', level=1)

    # Confusion matrices info
    cm_intro = doc.add_paragraph()
    cm_intro.add_run(
        f"The following confusion matrices show normalized classification performance for the "
        f"best-performing model per task selected by {selection_metric_label}. "
        f"Values represent the proportion of each true class classified into each predicted class."
    )
    doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # DEFINITE CONFUSION MATRICES (Figures 8-10)
    # ---------------------------------------------------------------------------

    doc.add_heading('Definite Cognitive Status', level=2)

    # Figure 8: Def NCD vs Dementia
    fig8_path = os.path.join(thesis_output_dir, "Figure_8_Definite_Confusion_Matrix_Task1.png")
    if os.path.exists(fig8_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 8. ").bold = True
        caption.add_run(format_cm_caption(
            "1. Def Normal vs. Def Dementia",
            "Confusion matrix: Definite NCD vs. Dementia"
        ))

        doc.add_picture(fig8_path, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("Figure 8 not available.")

    doc.add_paragraph()

    # Figure 9: Def NCD vs MCI
    fig9_path = os.path.join(thesis_output_dir, "Figure_9_Definite_Confusion_Matrix_Task3.png")
    if os.path.exists(fig9_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 9. ").bold = True
        caption.add_run(format_cm_caption(
            "3. Def Normal vs. Def MCI",
            "Confusion matrix: Definite NCD vs. MCI"
        ))

        doc.add_picture(fig9_path, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("Figure 9 not available.")

    doc.add_paragraph()

    # Figure 10: Def Multiclass
    fig10_path = os.path.join(thesis_output_dir, "Figure_10_Definite_Confusion_Matrix_Task5.png")
    if os.path.exists(fig10_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 10. ").bold = True
        caption.add_run(format_cm_caption(
            "5. Multi: Def Normal vs. Def MCI vs. Def Dementia",
            "Confusion matrix: Definite three-way classification"
        ))

        doc.add_picture(fig10_path, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("Figure 10 not available.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # DEFINITE + POSSIBLE CONFUSION MATRICES (Figures 11-13)
    # ---------------------------------------------------------------------------

    doc.add_heading('Definite + Possible Cognitive Status', level=2)

    # Figure 11: Def+Pos NCD vs Dementia
    fig11_path = os.path.join(thesis_output_dir, "Figure_11_DefPos_Confusion_Matrix_Task2.png")
    if os.path.exists(fig11_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 11. ").bold = True
        caption.add_run(format_cm_caption(
            "2. Def+Pos Normal vs. Def+Pos Dementia",
            "Confusion matrix: Def+Pos NCD vs. Dementia"
        ))

        doc.add_picture(fig11_path, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("Figure 11 not available.")

    doc.add_paragraph()

    # Figure 12: Def+Pos NCD vs MCI
    fig12_path = os.path.join(thesis_output_dir, "Figure_12_DefPos_Confusion_Matrix_Task4.png")
    if os.path.exists(fig12_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 12. ").bold = True
        caption.add_run(format_cm_caption(
            "4. Def+Pos Normal vs. Def+Pos MCI",
            "Confusion matrix: Def+Pos NCD vs. MCI"
        ))

        doc.add_picture(fig12_path, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("Figure 12 not available.")

    doc.add_paragraph()

    # Figure 13: Def+Pos Multiclass
    fig13_path = os.path.join(thesis_output_dir, "Figure_13_DefPos_Confusion_Matrix_Task6.png")
    if os.path.exists(fig13_path):
        caption = doc.add_paragraph(style='TableCaption')
        caption.add_run("Figure 13. ").bold = True
        caption.add_run(format_cm_caption(
            "6. Multi: Def+Pos Normal vs. Def+Pos MCI vs. Def+Pos Dementia",
            "Confusion matrix: Def+Pos three-way classification"
        ))

        doc.add_picture(fig13_path, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("Figure 13 not available.")

    doc.add_page_break()

    # ===========================================================================
    # SUPPLEMENTARY INFORMATION
    # ===========================================================================

    doc.add_heading('Supplementary Information', level=1)

    supp = doc.add_paragraph()
    supp.add_run("Methods Summary").bold = True
    doc.add_paragraph()

    methods_text = """Machine learning classification models were trained using linked administrative health data from DAD, NACRS, physician claims, and PIN. Six predefined tasks were analyzed: four binary tasks (definite and definite+possible NCD vs Dementia, and NCD vs MCI) and two three-class tasks (NCD vs MCI vs Dementia; definite and definite+possible).

Predictor preprocessing included median imputation for numeric fields, explicit 'missing' category handling for categorical fields, and one-hot encoding. To reduce label leakage and preserve a claims-based prediction setting, demographics (age, sex, race, education), cognitive test scores (MMSE/MoCA), outcome proxies, and identifiers were excluded from model features.

Each task used a single stratified 70/30 train-test split (random_state=42). Hyperparameter tuning used 3-fold StratifiedKFold on the training split. For binary tasks, probability thresholds were tuned via repeated stratified CV (5 folds x 3 repeats) to optimize balanced accuracy. Class imbalance handling used class weights and selective SMOTE inside CV/training pipelines only (not on held-out test data).

Primary reported discrimination metrics were AUC-ROC, PR-AUC, sensitivity, specificity, PPV, NPV, F1, and balanced accuracy. Uncertainty was quantified with 500 bootstrap resamples (BCa/percentile CI pipeline outputs). Pairwise model comparisons used DeLong tests (AUC) and McNemar tests (error rates), with Bonferroni and Benjamini-Hochberg FDR correction.

Calibration (slope/intercept, Brier score, Hosmer-Lemeshow), decision-curve analysis, class-imbalance reports, EPV/sample-size checks, and TRIPOD checklist outputs were generated for reproducibility support. The Jaakkimainen administrative algorithm was retained as the benchmark for dementia-vs-NCD tasks only. Results should be interpreted in the context of a high-risk referral cohort rather than population screening."""

    doc.add_paragraph(methods_text)

    task_summary_lines = build_task_design_summary(output_dir)
    if task_summary_lines:
        doc.add_paragraph()
        task_header = doc.add_paragraph()
        task_header.add_run("Task-Level Split and Imbalance Summary").bold = True
        for line in task_summary_lines:
            doc.add_paragraph(f"- {line}")

    stats_summary = build_statistical_test_summary(output_dir)
    if stats_summary:
        doc.add_paragraph()
        stats_header = doc.add_paragraph()
        stats_header.add_run("Pairwise Statistical Test Summary").bold = True
        for test_name in ["DeLong", "McNemar"]:
            info = stats_summary.get(test_name)
            if not info:
                continue
            doc.add_paragraph(
                f"- {test_name}: files={info['files']}, tests={info['total_tests']}, "
                f"nominal p<0.05={info['nominal_sig']}, FDR-significant={info['fdr_sig']}"
            )

    doc.add_paragraph()
    ref_title = doc.add_paragraph()
    ref_title.add_run("Reference").bold = True

    # Vancouver style reference
    ref = doc.add_paragraph()
    ref.add_run("1. Jaakkimainen RL, Bronskill SE, Tierney MC, et al. Identification of Physician-Diagnosed Alzheimer's Disease and Related Dementias in Population-Based Administrative Data: A Validation Study Using Family Physicians' Electronic Medical Records. ")
    ref.add_run("J Alzheimers Dis. ").italic = True
    ref.add_run("2016;54(1):337-349. doi:10.3233/JAD-160105")

    # ===========================================================================
    # SAVE DOCUMENT
    # ===========================================================================

    doc_path = os.path.join(thesis_output_dir, "Thesis_Tables_Figures_Publication.docx")
    doc.save(doc_path)

    print(f"\n{'='*50}")
    print("WORD DOCUMENT GENERATED SUCCESSFULLY")
    print(f"{'='*50}")
    print(f"Saved to: {doc_path}")
    print(f"File size: {os.path.getsize(doc_path):,} bytes")

    return doc_path


# ==============================================================================
# CALLABLE FUNCTION FOR INTEGRATION WITH original_parallel.py
# ==============================================================================

def run_thesis_outputs(data_df, results_list, output_base_dir):
    """
    Generate all thesis tables and figures from analysis results.

    This function is designed to be called from original_parallel.py at the end
    of the analysis pipeline.

    Parameters:
    -----------
    data_df : DataFrame
        Original data with cognitive_status column
    results_list : list
        List of result dictionaries from model evaluation
    output_base_dir : str
        Base output directory (e.g., cognitive_analysis_output_PARALLEL)

    Returns:
    --------
    dict : Paths to all generated outputs
    """
    print("\n" + "=" * 70)
    print("GENERATING THESIS TABLES AND FIGURES")
    print("=" * 70)

    # Setup output directory
    thesis_output = os.path.join(output_base_dir, "thesis_outputs")
    os.makedirs(thesis_output, exist_ok=True)

    generated_files = {}

    # Prepare data
    df = data_df.copy()
    df['cognitive_status'] = df['cognitive_status'].fillna('unknown').astype(str)
    df['cognitive_status'] = df['cognitive_status'].str.lower().str.strip().str.replace(' ', '_', regex=False)

    # =========================================================================
    # TABLE 1: DEFINITE Cognitive Status
    # =========================================================================
    try:
        definite_groups = {
            'Definite NCD': ['definite_normal'],
            'Definite MCI': ['definite_mci'],
            'Definite Dementia': ['definite_dementia']
        }

        table1, _ = generate_participant_characteristics_table(
            df, definite_groups,
            "Table_1_Participant_Characteristics_Definite",
            thesis_output
        )
        generated_files['table_1'] = os.path.join(thesis_output, "Table_1_Participant_Characteristics_Definite.csv")
    except Exception as e:
        print(f"  WARNING: Table 1 generation failed: {e}")

    # =========================================================================
    # TABLE 1b: DEFINITE + POSSIBLE Cognitive Status
    # =========================================================================
    try:
        defpos_groups = {
            'Def+Pos NCD': ['definite_normal', 'possible_normal'],
            'Def+Pos MCI': ['definite_mci', 'possible_mci'],
            'Def+Pos Dementia': ['definite_dementia', 'possible_dementia']
        }

        table1b, _ = generate_participant_characteristics_table(
            df, defpos_groups,
            "Table_1b_Participant_Characteristics_DefPos",
            thesis_output
        )
        generated_files['table_1b'] = os.path.join(thesis_output, "Table_1b_Participant_Characteristics_DefPos.csv")
    except Exception as e:
        print(f"  WARNING: Table 1b generation failed: {e}")

    # =========================================================================
    # Load/Create results DataFrame for Tables 3-4 and Figures
    # =========================================================================
    # Try to load from publication_table.csv first (has formatted CIs)
    results_df = load_existing_results(output_base_dir)

    if results_df is None:
        print("  NOTE: Could not load publication table - using raw results")
        # Would need to format results_list into proper DataFrame here
        return generated_files

    # Select confusion-matrix source models by AUC-ROC to align with
    # publication_visualizations.generate_confusion_matrices output selection.
    selection_metric_label = "AUC-ROC"
    best_models_by_task = build_best_models_by_task(
        results_df,
        selection_metric_col='AUC (95% CI)'
    )

    # =========================================================================
    # SECTION A: DEFINITE COGNITIVE STATUS TABLES
    # =========================================================================

    # TABLE 2: Definite NCD vs. Dementia (with Jaakkimainen reference)
    try:
        generate_performance_metrics_table(
            results_df, ['1. Def Normal vs. Def Dementia'],
            "Table_2_Definite_NCD_vs_Dementia",
            thesis_output,
            include_jaakkimainen=True
        )
        generated_files['table_2_def_dementia'] = os.path.join(thesis_output, "Table_2_Definite_NCD_vs_Dementia.csv")
    except Exception as e:
        print(f"  WARNING: Table 2 (Definite NCD vs Dementia) generation failed: {e}")

    # TABLE 3: Definite NCD vs. MCI
    try:
        generate_performance_metrics_table(
            results_df, ['3. Def Normal vs. Def MCI'],
            "Table_3_Definite_NCD_vs_MCI",
            thesis_output,
            include_jaakkimainen=False
        )
        generated_files['table_3_def_mci'] = os.path.join(thesis_output, "Table_3_Definite_NCD_vs_MCI.csv")
    except Exception as e:
        print(f"  WARNING: Table 3 (Definite NCD vs MCI) generation failed: {e}")

    # TABLE 4: Definite Multiclass Classification
    try:
        generate_performance_metrics_table(
            results_df, ['5. Multi: Def Normal vs. Def MCI vs. Def Dementia'],
            "Table_4_Definite_Multiclass",
            thesis_output,
            include_jaakkimainen=False
        )
        generated_files['table_4_def_multi'] = os.path.join(thesis_output, "Table_4_Definite_Multiclass.csv")
    except Exception as e:
        print(f"  WARNING: Table 4 (Definite Multiclass) generation failed: {e}")

    # =========================================================================
    # SECTION A: DEFINITE COGNITIVE STATUS FIGURES
    # =========================================================================

    # FIGURE 1: Definite NCD vs Dementia (with Jaakkimainen)
    try:
        generate_thesis_forest_plot(
            results_df,
            '1. Def Normal vs. Def Dementia',
            'Figure_1_Definite_NCD_vs_Dementia',
            thesis_output,
            include_jaakkimainen=True,
            metric='balanced_accuracy'
        )
        generated_files['figure_1'] = os.path.join(thesis_output, "Figure_1_Definite_NCD_vs_Dementia.png")
    except Exception as e:
        print(f"  WARNING: Figure 1 generation failed: {e}")

    # FIGURE 2: Definite NCD vs MCI
    try:
        generate_thesis_forest_plot(
            results_df,
            '3. Def Normal vs. Def MCI',
            'Figure_2_Definite_NCD_vs_MCI',
            thesis_output,
            include_jaakkimainen=False,
            metric='balanced_accuracy'
        )
        generated_files['figure_2'] = os.path.join(thesis_output, "Figure_2_Definite_NCD_vs_MCI.png")
    except Exception as e:
        print(f"  WARNING: Figure 2 generation failed: {e}")

    # FIGURE 3: Definite Multiclass
    try:
        generate_thesis_forest_plot(
            results_df,
            '5. Multi: Def Normal vs. Def MCI vs. Def Dementia',
            'Figure_3_Definite_Multiclass',
            thesis_output,
            include_jaakkimainen=False,
            metric='balanced_accuracy'
        )
        generated_files['figure_3'] = os.path.join(thesis_output, "Figure_3_Definite_Multiclass.png")
    except Exception as e:
        print(f"  WARNING: Figure 3 generation failed: {e}")

    # =========================================================================
    # SECTION B: DEFINITE + POSSIBLE COGNITIVE STATUS TABLES
    # =========================================================================

    # TABLE 6: Def+Pos NCD vs. Dementia
    try:
        generate_performance_metrics_table(
            results_df, ['2. Def+Pos Normal vs. Def+Pos Dementia'],
            "Table_6_DefPos_NCD_vs_Dementia",
            thesis_output,
            include_jaakkimainen=False
        )
        generated_files['table_6_defpos_dementia'] = os.path.join(thesis_output, "Table_6_DefPos_NCD_vs_Dementia.csv")
    except Exception as e:
        print(f"  WARNING: Table 6 (Def+Pos NCD vs Dementia) generation failed: {e}")

    # TABLE 7: Def+Pos NCD vs. MCI
    try:
        generate_performance_metrics_table(
            results_df, ['4. Def+Pos Normal vs. Def+Pos MCI'],
            "Table_7_DefPos_NCD_vs_MCI",
            thesis_output,
            include_jaakkimainen=False
        )
        generated_files['table_7_defpos_mci'] = os.path.join(thesis_output, "Table_7_DefPos_NCD_vs_MCI.csv")
    except Exception as e:
        print(f"  WARNING: Table 7 (Def+Pos NCD vs MCI) generation failed: {e}")

    # TABLE 8: Def+Pos Multiclass Classification
    try:
        generate_performance_metrics_table(
            results_df, ['6. Multi: Def+Pos Normal vs. Def+Pos MCI vs. Def+Pos Dementia'],
            "Table_8_DefPos_Multiclass",
            thesis_output,
            include_jaakkimainen=False
        )
        generated_files['table_8_defpos_multi'] = os.path.join(thesis_output, "Table_8_DefPos_Multiclass.csv")
    except Exception as e:
        print(f"  WARNING: Table 8 (Def+Pos Multiclass) generation failed: {e}")

    # =========================================================================
    # SECTION B: DEFINITE + POSSIBLE COGNITIVE STATUS FIGURES
    # =========================================================================

    # FIGURE 4: Def+Pos NCD vs Dementia
    try:
        generate_thesis_forest_plot(
            results_df,
            '2. Def+Pos Normal vs. Def+Pos Dementia',
            'Figure_4_DefPos_NCD_vs_Dementia',
            thesis_output,
            include_jaakkimainen=False,
            metric='balanced_accuracy'
        )
        generated_files['figure_4'] = os.path.join(thesis_output, "Figure_4_DefPos_NCD_vs_Dementia.png")
    except Exception as e:
        print(f"  WARNING: Figure 4 generation failed: {e}")

    # FIGURE 5: Def+Pos NCD vs MCI
    try:
        generate_thesis_forest_plot(
            results_df,
            '4. Def+Pos Normal vs. Def+Pos MCI',
            'Figure_5_DefPos_NCD_vs_MCI',
            thesis_output,
            include_jaakkimainen=False,
            metric='balanced_accuracy'
        )
        generated_files['figure_5'] = os.path.join(thesis_output, "Figure_5_DefPos_NCD_vs_MCI.png")
    except Exception as e:
        print(f"  WARNING: Figure 5 generation failed: {e}")

    # FIGURE 6: Def+Pos Multiclass
    try:
        generate_thesis_forest_plot(
            results_df,
            '6. Multi: Def+Pos Normal vs. Def+Pos MCI vs. Def+Pos Dementia',
            'Figure_6_DefPos_Multiclass',
            thesis_output,
            include_jaakkimainen=False,
            metric='balanced_accuracy'
        )
        generated_files['figure_6'] = os.path.join(thesis_output, "Figure_6_DefPos_Multiclass.png")
    except Exception as e:
        print(f"  WARNING: Figure 6 generation failed: {e}")

    # =========================================================================
    # SECTION C: SUPPLEMENTARY - Feature Importance
    # =========================================================================

    # FIGURE 7: Aggregated Feature Importance
    try:
        stats_dir = os.path.join(output_base_dir, "statistical_comparisons")
        generate_aggregated_feature_importance_plot(
            stats_dir,
            thesis_output,
            top_n=10
        )
        generated_files['figure_7'] = os.path.join(thesis_output, "Figure_7_Feature_Importance.png")
    except Exception as e:
        print(f"  WARNING: Figure 7 generation failed: {e}")

    # =========================================================================
    # SECTION D: CONFUSION MATRICES FOR BEST PERFORMING MODELS
    # =========================================================================

    try:
        confusion_matrix_dir = os.path.join(output_base_dir, "publication_figures", "confusion_matrices")
        if os.path.exists(confusion_matrix_dir):
            print("\n=== Generating Confusion Matrices (Section D) ===")
            cm_files = generate_all_confusion_matrices(
                confusion_matrix_dir,
                thesis_output,
                best_models_by_task=best_models_by_task,
                selection_metric_label=selection_metric_label
            )
            for key, path in cm_files.items():
                generated_files[key] = path
        else:
            print(f"  NOTE: Confusion matrix directory not found: {confusion_matrix_dir}")
    except Exception as e:
        print(f"  WARNING: Confusion matrix generation failed: {e}")

    # =========================================================================
    # GENERATE UNIFIED WORD DOCUMENT
    # =========================================================================
    try:
        if DOCX_AVAILABLE:
            doc_path = generate_thesis_word_document(
                output_base_dir,
                thesis_output,
                best_models_by_task=best_models_by_task,
                selection_metric_label=selection_metric_label
            )
            if doc_path:
                generated_files['word_document'] = doc_path
        else:
            print("  NOTE: python-docx not available - Word document not generated")
    except Exception as e:
        print(f"  WARNING: Word document generation failed: {e}")

    # Print summary
    print("\n" + "-" * 50)
    print("THESIS OUTPUTS GENERATED:")
    print("-" * 50)
    for name, path in generated_files.items():
        if os.path.exists(path):
            size = os.path.getsize(path)
            print(f"  ✓ {name}: {os.path.basename(path)} ({size:,} bytes)")
        else:
            print(f"  ✗ {name}: FAILED")
    print(f"\nAll outputs in: {thesis_output}")

    return generated_files


# ==============================================================================
# MAIN EXECUTION (for standalone running)
# ==============================================================================

def main():
    """Generate all thesis tables and figures."""
    print("=" * 70)
    print("THESIS TABLES AND FIGURES GENERATOR")
    print("=" * 70)
    print(f"Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Setup paths
    script_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in dir() else os.getcwd()
    data_path = os.path.join(script_dir, DATA_PATH)
    output_base = os.path.join(script_dir, OUTPUT_BASE_DIR)
    # Keep CLI output location consistent with run_thesis_outputs().
    thesis_output = os.path.join(output_base, THESIS_OUTPUT_DIR)

    os.makedirs(thesis_output, exist_ok=True)

    # Check for data file
    if not os.path.exists(data_path):
        print(f"ERROR: Data file not found: {data_path}")
        return

    # Load data
    print("\n=== Loading Data ===")
    df = pd.read_csv(data_path)
    print(f"Loaded: {df.shape[0]} rows, {df.shape[1]} columns")

    # Prepare cognitive status
    df['cognitive_status'] = df['cognitive_status'].fillna('unknown').astype(str)
    df['cognitive_status'] = df['cognitive_status'].str.lower().str.strip().str.replace(' ', '_', regex=False)

    print("\nCognitive status distribution:")
    print(df['cognitive_status'].value_counts())

    # =========================================================================
    # TABLE 1: DEFINITE Cognitive Status
    # =========================================================================
    definite_groups = {
        'Definite NCD': ['definite_normal'],
        'Definite MCI': ['definite_mci'],
        'Definite Dementia': ['definite_dementia']
    }

    generate_participant_characteristics_table(
        df,
        definite_groups,
        "Table_1_Participant_Characteristics_Definite",
        thesis_output
    )

    # =========================================================================
    # TABLE 1b: DEFINITE + POSSIBLE Cognitive Status
    # =========================================================================
    defpos_groups = {
        'Def+Pos NCD': ['definite_normal', 'possible_normal'],
        'Def+Pos MCI': ['definite_mci', 'possible_mci'],
        'Def+Pos Dementia': ['definite_dementia', 'possible_dementia']
    }

    generate_participant_characteristics_table(
        df,
        defpos_groups,
        "Table_1b_Participant_Characteristics_DefPos",
        thesis_output
    )

    # =========================================================================
    # Load existing results for Tables 3-4 and Figures
    # =========================================================================
    print("\n=== Loading Existing Analysis Results ===")
    results_df = load_existing_results(output_base)

    if results_df is None:
        print("ERROR: Could not load existing results. Run original_parallel.py first.")
        return

    print(f"Loaded results: {len(results_df)} model evaluations")

    selection_metric_label = "AUC-ROC"
    best_models_by_task = build_best_models_by_task(
        results_df,
        selection_metric_col='AUC (95% CI)'
    )

    # =========================================================================
    # SECTION A: DEFINITE COGNITIVE STATUS TABLES
    # =========================================================================

    # TABLE 2: Definite NCD vs Dementia
    generate_performance_metrics_table(
        results_df,
        ['1. Def Normal vs. Def Dementia'],
        "Table_2_Definite_NCD_vs_Dementia",
        thesis_output,
        include_jaakkimainen=True
    )

    # TABLE 3: Definite NCD vs MCI
    generate_performance_metrics_table(
        results_df,
        ['3. Def Normal vs. Def MCI'],
        "Table_3_Definite_NCD_vs_MCI",
        thesis_output,
        include_jaakkimainen=False
    )

    # TABLE 4: Definite Multiclass Classification Performance
    generate_performance_metrics_table(
        results_df,
        ['5. Multi: Def Normal vs. Def MCI vs. Def Dementia'],
        "Table_4_Definite_Multiclass",
        thesis_output,
        include_jaakkimainen=False
    )

    # =========================================================================
    # SECTION B: DEFINITE + POSSIBLE COGNITIVE STATUS TABLES
    # =========================================================================

    # TABLE 6: Def+Pos NCD vs Dementia
    generate_performance_metrics_table(
        results_df,
        ['2. Def+Pos Normal vs. Def+Pos Dementia'],
        "Table_6_DefPos_NCD_vs_Dementia",
        thesis_output,
        include_jaakkimainen=False
    )

    # TABLE 7: Def+Pos NCD vs MCI
    generate_performance_metrics_table(
        results_df,
        ['4. Def+Pos Normal vs. Def+Pos MCI'],
        "Table_7_DefPos_NCD_vs_MCI",
        thesis_output,
        include_jaakkimainen=False
    )

    # TABLE 8: Def+Pos Multiclass Classification Performance
    generate_performance_metrics_table(
        results_df,
        ['6. Multi: Def+Pos Normal vs. Def+Pos MCI vs. Def+Pos Dementia'],
        "Table_8_DefPos_Multiclass",
        thesis_output,
        include_jaakkimainen=False
    )

    # =========================================================================
    # SECTION A FIGURES (DEFINITE COGNITIVE STATUS)
    # =========================================================================

    # FIGURE 1: Forest Plot - Definite NCD vs Dementia (with Jaakkimainen)
    generate_thesis_forest_plot(
        results_df,
        '1. Def Normal vs. Def Dementia',
        'Figure_1_Definite_NCD_vs_Dementia',
        thesis_output,
        include_jaakkimainen=True,
        metric='balanced_accuracy'
    )

    # FIGURE 2: Forest Plot - Definite NCD vs MCI
    generate_thesis_forest_plot(
        results_df,
        '3. Def Normal vs. Def MCI',
        'Figure_2_Definite_NCD_vs_MCI',
        thesis_output,
        include_jaakkimainen=False,
        metric='balanced_accuracy'
    )

    # FIGURE 3: Forest Plot - Definite Multiclass
    generate_thesis_forest_plot(
        results_df,
        '5. Multi: Def Normal vs. Def MCI vs. Def Dementia',
        'Figure_3_Definite_Multiclass',
        thesis_output,
        include_jaakkimainen=False,
        metric='balanced_accuracy'
    )

    # =========================================================================
    # SECTION B FIGURES (DEFINITE + POSSIBLE COGNITIVE STATUS)
    # =========================================================================

    # FIGURE 4: Forest Plot - Def+Pos NCD vs Dementia
    generate_thesis_forest_plot(
        results_df,
        '2. Def+Pos Normal vs. Def+Pos Dementia',
        'Figure_4_DefPos_NCD_vs_Dementia',
        thesis_output,
        include_jaakkimainen=False,
        metric='balanced_accuracy'
    )

    # FIGURE 5: Forest Plot - Def+Pos NCD vs MCI
    generate_thesis_forest_plot(
        results_df,
        '4. Def+Pos Normal vs. Def+Pos MCI',
        'Figure_5_DefPos_NCD_vs_MCI',
        thesis_output,
        include_jaakkimainen=False,
        metric='balanced_accuracy'
    )

    # FIGURE 6: Forest Plot - Def+Pos Multiclass
    generate_thesis_forest_plot(
        results_df,
        '6. Multi: Def+Pos Normal vs. Def+Pos MCI vs. Def+Pos Dementia',
        'Figure_6_DefPos_Multiclass',
        thesis_output,
        include_jaakkimainen=False,
        metric='balanced_accuracy'
    )

    # =========================================================================
    # SECTION C: FIGURE 7 - Aggregated Feature Importance
    # =========================================================================
    stats_dir = os.path.join(output_base, "statistical_comparisons")
    generate_aggregated_feature_importance_plot(
        stats_dir,
        thesis_output,
        top_n=10
    )

    # =========================================================================
    # SECTION D: FIGURES 8-13 - Confusion Matrices for Best Models
    # =========================================================================
    confusion_matrix_dir = os.path.join(output_base, "publication_figures", "confusion_matrices")
    if os.path.exists(confusion_matrix_dir):
        generate_all_confusion_matrices(
            confusion_matrix_dir,
            thesis_output,
            best_models_by_task=best_models_by_task,
            selection_metric_label=selection_metric_label
        )
    else:
        print(f"\nWARNING: Confusion matrix directory not found: {confusion_matrix_dir}")

    # =========================================================================
    # GENERATE UNIFIED WORD DOCUMENT
    # =========================================================================
    if DOCX_AVAILABLE:
        generate_thesis_word_document(
            output_base,
            thesis_output,
            best_models_by_task=best_models_by_task,
            selection_metric_label=selection_metric_label
        )
    else:
        print("\nWARNING: python-docx not available. Install with: pip install python-docx")

    # =========================================================================
    # Summary
    # =========================================================================
    print("\n" + "=" * 70)
    print("THESIS TABLES AND FIGURES GENERATION COMPLETE")
    print("=" * 70)
    print(f"\nAll outputs saved to: {thesis_output}")
    print("\nGenerated files:")

    for f in sorted(os.listdir(thesis_output)):
        filepath = os.path.join(thesis_output, f)
        size = os.path.getsize(filepath)
        print(f"  - {f} ({size:,} bytes)")

    print(f"\nCompleted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


if __name__ == "__main__":
    main()
# fmt: on
